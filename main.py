from post import get_stu_task
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
import os
from urllib.parse import urljoin
from post import get_stu_task
from docx.shared import Pt, Cm
import re
from utils import download_image
from docx.shared import Pt, Cm, Inches

# 预备操作
img_dir = "temp_images"
os.makedirs(img_dir, exist_ok=True)


# ============================================================================
# 常量设置
# 循环

def main(i = 1):
    # task_id = "5465"#5461-5479
    cid = "284"
    SESSION_ID = ""  # 替换实际PHPSESSID
    BASE_URL = "http://10.200.64.39"

    # ===============================================================================
    # 爬取
    task_id = i+5465 if i<=14 else i+5446 #5461-5479
    res = get_stu_task(task_id, cid, SESSION_ID, BASE_URL)

    data_r = {
        "实验室名称": "商务数据分析王刚老师上机实验室",
        "实验项目名称": res['datas']['taskName'],
        "实验学时": "2学时", 
        "实验原理": "",
        "实验目的": "",
        "实验内容": res['datas']['taskName'], 
        "实验环境": "电脑",
        "实验步骤": "",
        "实验结果及分析": "详见步骤中的运行结果截图",
        "实验结论": "",
        "总结及心得体会":"人工智能飞速发展，平台的实验真的都让我学到了很多有用的知识"
    }
    # ===============================================================================

    # 获取文件名称
    OUTPUT_FILE = './doc2/'+res['datas']['taskName']+".docx"
    # ===============================================================================
    # 清洗
    guide_html = res['datas']['reference']  # 使用reference字段包含完整步骤
    soup = BeautifulSoup(guide_html, 'html.parser')

    k="实验目的"
    data = {}
    for element in soup.find_all(True):
        # strong里的内容会被双次捕获，这里直接去除了
        if element.name != 'p':
            continue
        # 获取标签内容
        text = element.get_text().strip()
        # 如果有图片
        img = element.find('img')
        if img:
            # img_uri = 
            # img_url = urljoin(BASE_URL, img['src'])
            # img_name = download_image(img_url, img_dir, SESSION_ID)
            # 原本去获取网络图片，现在直接拿我做的图
            data.setdefault(k,[]).append(img)
            continue
        # 没有图片就处理文字
        if text == '':
            continue
        if re.search(r'[【】]', text):
            k = re.sub(r'[【】]', '', text).strip()
            data.setdefault(k,[])
        else:
            data[k].append(text)
    # ===============================================================================
    # 创建文件配置
    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal'].font.size = Pt(15)
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # ===============================================================================
    num = [
    "一",  
    "二",
    "三",
    "四",
    "五",
    "六",
    "七",
    "八",
    "九",
    "十",
    "十一"
    ]
    _=0
    for key, value in data_r.items():
        p = doc.add_paragraph()
        p.add_run(num[_]+'、'+key).bold = True

        
        if key == "实验结论":
            key = "实验原理"
        try:
            if data[key]:
                tem = 1
                for temp in data[key]:
                    try:
                        if temp.name =='img':
                            # "c:\Users\Administrator\Desktop\sb\1\1.png"
                            path = os.path.join("c:/Users/Administrator/Desktop/sb/", str(i)+"/"+str(tem)+".png")
                            if os.path.exists(path):
                                try:
                                    doc.add_picture(path, width=Cm(15))
                                except Exception as e:
                                    print(e)
                                tem +=1
                            else:
                                print(path+" is not exist!!!!")
                    except:
                        p = doc.add_paragraph()
                        p.add_run(temp)
                        
        except:
            p = doc.add_paragraph()
            p.add_run(value)
        finally:
            _+=1

    #完成所有处理
    doc.save(OUTPUT_FILE)


if __name__ == "__main__":
    # for i in range(1,20):
    #     main(i)
    main(19)